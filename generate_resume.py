#!/usr/bin/env python3
"""Generate Ananya Rai Paul's Data Analyst resume (PDF + DOCX)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor, Twips
from reportlab.lib.colors import HexColor, black
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    HRFlowable,
    KeepTogether,
)

OUT_DIR = Path("/workspace")
PDF_PATH = OUT_DIR / "Ananya_Rai_Paul_Data_Analyst_Resume.pdf"
DOCX_PATH = OUT_DIR / "Ananya_Rai_Paul_Data_Analyst_Resume.docx"

NAVY = HexColor("#1a1a2e")
GRAY = HexColor("#333333")
MUTED = HexColor("#444444")


def build_pdf() -> None:
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=0.6 * inch,
        rightMargin=0.6 * inch,
        topMargin=0.35 * inch,
        bottomMargin=0.35 * inch,
    )

    styles = {
        "name": ParagraphStyle(
            "Name",
            fontName="Helvetica-Bold",
            fontSize=15,
            textColor=NAVY,
            alignment=1,
            spaceAfter=2,
            leading=17,
        ),
        "contact": ParagraphStyle(
            "Contact",
            fontName="Helvetica",
            fontSize=9,
            textColor=GRAY,
            alignment=1,
            spaceAfter=4,
            leading=11,
        ),
        "section": ParagraphStyle(
            "Section",
            fontName="Helvetica-Bold",
            fontSize=10.5,
            textColor=NAVY,
            spaceBefore=6,
            spaceAfter=2,
            leading=12,
        ),
        "body": ParagraphStyle(
            "Body",
            fontName="Helvetica",
            fontSize=9,
            textColor=GRAY,
            leading=11.5,
            spaceAfter=1,
            alignment=0,
        ),
        "bullet": ParagraphStyle(
            "Bullet",
            fontName="Helvetica",
            fontSize=9,
            textColor=GRAY,
            leading=11.5,
            leftIndent=10,
            spaceAfter=0.5,
        ),
        "job_title": ParagraphStyle(
            "JobTitle",
            fontName="Helvetica-Bold",
            fontSize=9.5,
            textColor=black,
            spaceBefore=3,
            spaceAfter=0.5,
            leading=11,
        ),
        "edu_title": ParagraphStyle(
            "EduTitle",
            fontName="Helvetica-Bold",
            fontSize=9.5,
            textColor=black,
            spaceBefore=2,
            spaceAfter=0.5,
            leading=11,
        ),
        "edu_sub": ParagraphStyle(
            "EduSub",
            fontName="Helvetica",
            fontSize=9,
            textColor=GRAY,
            leading=11,
            spaceAfter=0.5,
        ),
    }

    story = []

    story.append(Paragraph("ANANYA RAI PAUL", styles["name"]))
    story.append(
        Paragraph(
            "anannyaraipaul@gmail.com | linkedin.com/in/annyarp | github.com/anannyarp12",
            styles["contact"],
        )
    )
    story.append(
        HRFlowable(width="100%", thickness=1, color=NAVY, spaceBefore=0, spaceAfter=3)
    )

    story.append(Paragraph("SUMMARY", styles["section"]))
    story.append(
        HRFlowable(width="100%", thickness=0.5, color=MUTED, spaceBefore=0, spaceAfter=2)
    )
    story.append(
        Paragraph(
            "B.Sc. Statistics (Honours) graduate with hands-on experience turning large datasets "
            "into actionable business insight using SQL, Excel, and Power BI — including revenue "
            "trends, customer behavior patterns, and performance dashboards. Interested in applying "
            "analytical rigor to growth and marketing problems: understanding user behavior, "
            "campaign performance, and what drives retention and monetization.",
            styles["body"],
        )
    )

    story.append(Paragraph("TECHNICAL SKILLS", styles["section"]))
    story.append(
        HRFlowable(width="100%", thickness=0.5, color=MUTED, spaceBefore=0, spaceAfter=2)
    )
    skills = [
        "Query &amp; Data Handling: SQL (Joins, Aggregations, Filtering, Grouping)",
        "Spreadsheet Analysis: Microsoft Excel (Pivot Tables, VLOOKUP, Charts, Data Cleaning)",
        "Business Intelligence: Power BI (Interactive Dashboards, Data Modeling, DAX Basics)",
        "Programming: Python (working knowledge — scripting, data cleaning, basic analysis)",
        "Statistical Concepts: Hypothesis Testing, Probability, Regression, Exploratory Data Analysis",
        "Core Strengths: Analytical Thinking, Attention to Detail, Communication",
    ]
    for s in skills:
        story.append(Paragraph(f"• {s}", styles["bullet"]))

    story.append(Paragraph("PROJECTS", styles["section"]))
    story.append(
        HRFlowable(width="100%", thickness=0.5, color=MUTED, spaceBefore=0, spaceAfter=2)
    )

    projects = [
        (
            "Revenue &amp; Product Performance Analysis — SQL, Power BI",
            [
                "Queried and joined 2,000+ multi-region sales records using SQL (joins, aggregations, grouping) to surface revenue and product performance trends",
                "Identified consistently top-performing product categories and regional revenue gaps, informing prioritization discussions",
                "Built an interactive Power BI dashboard translating findings into stakeholder-ready KPIs; findings directly informed which product categories and regions merited increased marketing investment",
            ],
        ),
        (
            "Interactive Sales Performance Dashboard — Power BI",
            [
                "Designed a multi-view Power BI dashboard with drill-through pages to monitor revenue and product performance in real time",
                "Applied data modeling and DAX measures to compare performance across regions and time periods — similar to tracking acquisition and retention metrics over time",
                "Structured views for quick review, reducing time spent compiling manual reports",
            ],
        ),
        (
            "Retail Transaction &amp; Seasonal Trend Analysis — Microsoft Excel",
            [
                "Cleaned and analyzed 500,000+ retail transaction records using pivot tables and charts to identify key revenue contributors",
                "Uncovered seasonal sales spikes and recurring customer purchasing patterns through exploratory data analysis",
                "Identified peak-demand windows that could inform the timing of promotional and acquisition campaigns",
            ],
        ),
        (
            "Airline Delay Pattern Analysis — SQL, Excel, Power BI",
            [
                "Queried flight delay data across routes, airports, and time periods in SQL to detect recurring operational bottlenecks",
                "Combined SQL and Excel analysis to isolate the routes and time windows with the highest delay concentration",
                "Built a Power BI dashboard visualizing delay trends, supporting data-driven decision-making",
            ],
        ),
    ]

    for title, bullets in projects:
        block = [Paragraph(title, styles["job_title"])]
        for b in bullets:
            block.append(Paragraph(f"• {b}", styles["bullet"]))
        story.append(KeepTogether(block))

    story.append(Paragraph("EDUCATION", styles["section"]))
    story.append(
        HRFlowable(width="100%", thickness=0.5, color=MUTED, spaceBefore=0, spaceAfter=2)
    )
    story.append(Paragraph("B.Sc. Statistics (Honours)", styles["edu_title"]))
    story.append(
        Paragraph("University of North Bengal | CGPA: 7.70 / 10", styles["edu_sub"])
    )
    story.append(
        Paragraph(
            "Coursework: Probability, Regression, Hypothesis Testing, Multivariate Analysis, "
            "Time Series, Design of Experiments, Statistical Quality Control, Exploratory Data Analysis",
            styles["edu_sub"],
        )
    )
    story.append(Spacer(1, 2))
    story.append(
        Paragraph(
            "St. Joseph's High School — ISC Class XII (2021)",
            styles["edu_title"],
        )
    )
    story.append(
        Paragraph(
            "English, Mathematics, Physics, Chemistry | Percentage: 74%",
            styles["edu_sub"],
        )
    )

    story.append(Paragraph("CERTIFICATIONS &amp; LEARNING", styles["section"]))
    story.append(
        HRFlowable(width="100%", thickness=0.5, color=MUTED, spaceBefore=0, spaceAfter=2)
    )
    story.append(Paragraph("• IBM – Introduction to Data Analytics", styles["bullet"]))
    story.append(
        Paragraph("• Google – Fundamentals of Digital Marketing", styles["bullet"])
    )

    doc.build(story)


def set_run_font(run, name="Calibri", size=10, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A1A2E")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_docx() -> None:
    document = Document()
    for section in document.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    pf = style.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.08

    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name.add_run("ANANYA RAI PAUL")
    set_run_font(r, size=16, bold=True, color=(26, 26, 46))
    name.paragraph_format.space_after = Pt(2)

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = contact.add_run(
        "anannyaraipaul@gmail.com | linkedin.com/in/annyarp | github.com/anannyarp12"
    )
    set_run_font(r, size=10, color=(51, 51, 51))
    contact.paragraph_format.space_after = Pt(4)
    add_horizontal_line(contact)

    def section_heading(text: str):
        p = document.add_paragraph()
        r = p.add_run(text)
        set_run_font(r, size=11, bold=True, color=(26, 26, 46))
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        add_horizontal_line(p)
        return p

    def body(text: str):
        p = document.add_paragraph()
        r = p.add_run(text)
        set_run_font(r, size=10, color=(51, 51, 51))
        p.paragraph_format.space_after = Pt(3)
        return p

    def bullet(text: str):
        p = document.add_paragraph(style="List Bullet")
        p.clear()
        r = p.add_run(text)
        set_run_font(r, size=10, color=(51, 51, 51))
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(1)
        return p

    def bold_line(text: str):
        p = document.add_paragraph()
        r = p.add_run(text)
        set_run_font(r, size=10, bold=True, color=(0, 0, 0))
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        return p

    section_heading("SUMMARY")
    body(
        "B.Sc. Statistics (Honours) graduate with hands-on experience turning large datasets "
        "into actionable business insight using SQL, Excel, and Power BI — including revenue "
        "trends, customer behavior patterns, and performance dashboards. Interested in applying "
        "analytical rigor to growth and marketing problems: understanding user behavior, "
        "campaign performance, and what drives retention and monetization."
    )

    section_heading("TECHNICAL SKILLS")
    for s in [
        "Query & Data Handling: SQL (Joins, Aggregations, Filtering, Grouping)",
        "Spreadsheet Analysis: Microsoft Excel (Pivot Tables, VLOOKUP, Charts, Data Cleaning)",
        "Business Intelligence: Power BI (Interactive Dashboards, Data Modeling, DAX Basics)",
        "Programming: Python (working knowledge — scripting, data cleaning, basic analysis)",
        "Statistical Concepts: Hypothesis Testing, Probability, Regression, Exploratory Data Analysis",
        "Core Strengths: Analytical Thinking, Attention to Detail, Communication",
    ]:
        bullet(s)

    section_heading("PROJECTS")
    projects = [
        (
            "Revenue & Product Performance Analysis — SQL, Power BI",
            [
                "Queried and joined 2,000+ multi-region sales records using SQL (joins, aggregations, grouping) to surface revenue and product performance trends",
                "Identified consistently top-performing product categories and regional revenue gaps, informing prioritization discussions",
                "Built an interactive Power BI dashboard translating findings into stakeholder-ready KPIs; findings directly informed which product categories and regions merited increased marketing investment",
            ],
        ),
        (
            "Interactive Sales Performance Dashboard — Power BI",
            [
                "Designed a multi-view Power BI dashboard with drill-through pages to monitor revenue and product performance in real time",
                "Applied data modeling and DAX measures to compare performance across regions and time periods — similar to tracking acquisition and retention metrics over time",
                "Structured views for quick review, reducing time spent compiling manual reports",
            ],
        ),
        (
            "Retail Transaction & Seasonal Trend Analysis — Microsoft Excel",
            [
                "Cleaned and analyzed 500,000+ retail transaction records using pivot tables and charts to identify key revenue contributors",
                "Uncovered seasonal sales spikes and recurring customer purchasing patterns through exploratory data analysis",
                "Identified peak-demand windows that could inform the timing of promotional and acquisition campaigns",
            ],
        ),
        (
            "Airline Delay Pattern Analysis — SQL, Excel, Power BI",
            [
                "Queried flight delay data across routes, airports, and time periods in SQL to detect recurring operational bottlenecks",
                "Combined SQL and Excel analysis to isolate the routes and time windows with the highest delay concentration",
                "Built a Power BI dashboard visualizing delay trends, supporting data-driven decision-making",
            ],
        ),
    ]
    for title, bullets in projects:
        bold_line(title)
        for b in bullets:
            bullet(b)

    section_heading("EDUCATION")
    bold_line("B.Sc. Statistics (Honours)")
    body("University of North Bengal | CGPA: 7.70 / 10")
    body(
        "Coursework: Probability, Regression, Hypothesis Testing, Multivariate Analysis, "
        "Time Series, Design of Experiments, Statistical Quality Control, Exploratory Data Analysis"
    )
    bold_line("St. Joseph's High School — ISC Class XII (2021)")
    body("English, Mathematics, Physics, Chemistry | Percentage: 74%")

    section_heading("CERTIFICATIONS & LEARNING")
    bullet("IBM – Introduction to Data Analytics")
    bullet("Google – Fundamentals of Digital Marketing")

    document.save(str(DOCX_PATH))


if __name__ == "__main__":
    build_pdf()
    build_docx()
    print(f"Wrote {PDF_PATH}")
    print(f"Wrote {DOCX_PATH}")
